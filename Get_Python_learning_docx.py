# -*- coding=utf-8 -*-
'''
@author: Administrator
@time:2017年5月28日
@description:获取廖雪峰老师的学习资料,并保存为docx
'''

from bs4 import BeautifulSoup
import urllib2
import docx
import re

user_agent = 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/50.0.2661.102 Safari/537.36'
headers = {'User-Agent': user_agent}
document = docx.Document()
document.add_heading(u'廖雪峰-python教程2.7', 0)
def get_request_content(url):
    request = urllib2.Request(url, headers=headers)
    response = urllib2.urlopen(request)
    content = response.read().decode('utf-8')
    return content
def get_content(href):
    soup = BeautifulSoup(get_request_content(url_title+href), 'lxml')
    content_class='x-wiki-content'
    content=soup.find_all('div',content_class)
#     print 'content: ',content[0]
    return content[0]
url_index = 'http://www.liaoxuefeng.com/wiki/001374738125095c955c1e6d8bb493182103fac9270762a000'
soup = BeautifulSoup(get_request_content(url_index), 'lxml')
x=soup.find_all('div','x-sidebar-left-content')#获取章节
# print x
x=x[0]
list1 = x.find_all('a')
length=len(list1)#获取长度
index=1
url_title='http://www.liaoxuefeng.com'
for i in range(length):
    print 'index:',index
#     print list1[i]
    title=list1[i].string#章节标题
    print "title: ",title
    href = list1[i]['href']
    parent=list1[i].parent
#     print type(parent)
    bflag= parent.has_attr('style')
    if bflag:
#         print "设置成二级title"
        document.add_heading(title,2)
    else :
#         print "设置成一级title"
        document.add_heading(title,1)
    content=get_content(href)
    document.add_paragraph(unicode(content))
    index+=1
#     if index==3:break
document.save('demo.docx')
print 'done'
